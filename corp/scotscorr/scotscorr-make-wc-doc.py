#! /usr/bin/env python
# -*- coding: utf-8 -*-


"""Generate a list of letter word counts by gender, period and locality."""


import re
import codecs

from cStringIO import StringIO

# Package python-docx: https://python-docx.readthedocs.io/en/latest/
import docx
from docx.shared import Pt, Mm

import korpimport.util as korputil


class WordCountDocMakerError(Exception):

    pass


class WordCountDocMaker(korputil.InputProcessor):

    # Function for class attribute definition, not a method
    def _make_order(list_):
        return dict((key, num) for num, key in enumerate(list_))

    _genders = ['male', 'female', 'royal']
    _localities = [
        'Moray',
        'Invernessshire',
        'Sutherland',
        'Ross',
        'Aberdeenshire',
        'Angus',
        'Perthshire',
        'Lanarkshire',
        'Fife',
        'Lothian',
        'Borders',
        'Stirlingshire',
        'Ayrshire',
        'Argyllshire',
        'South-West',
        'Professional',
        'Court',
        'unlocalised',
    ]
    _gender_order = _make_order(_genders)
    _locality_order = _make_order(_localities)

    def __init__(self):
        super(WordCountDocMaker, self).__init__()
        self._letter_info = {}
        self._letter_order = []
        if self._opts.order_file:
            self._read_order_file()

    def _read_order_file(self):
        with codecs.open(self._opts.order_file, 'r',
                         encoding=self._input_encoding) as f:
            for line in f:
                if line.startswith('%FN:'):
                    mo = re.match(r'%\S+?:\s*(\S+?)\s*$', line)
                    filename = mo.group(1) if mo else ""
                    self._letter_order.append(filename)

    def process_input_stream(self, stream, filename=None):
        self._read_wordcount_file(stream)
        if not self._letter_order:
            self._compute_order()
        self._generate_output()
        self._write_output()

    def _read_wordcount_file(self, stream):
        for line in korputil.tsv_dictreader(stream):
            self._linenr += 1
            self._letter_info[line['fn']] = line

    def _compute_order(self):

        def make_letter_key(fn):
            info = self._letter_info[fn]
            return (self._gender_order[info['gender']],
                    info['period'],
                    self._locality_order[info['lcinf']],
                    fn)

        self._letter_order = self._letter_info.keys()
        self._letter_order.sort(key=make_letter_key)

    class OutputLetterData(object):

        def __init__(self, letter_data=None):
            self.gender = None
            self.period = None
            self.locality = None
            self.informant = None
            if letter_data:
                filename, self.gender, self.period, self.locality = tuple(
                    letter_data.get(key) for key in [
                        'fn', 'gender', 'period', 'lcinf'])
                mo = re.match(r'(.*?)\d{7}', filename)
                if not mo:
                    raise WordCountDocMakerError(
                        'Cannot find informant in filename: ' + filename)
                self.informant = mo.group(1)

    class LetterStats(object):

        def __init__(self):
            self.infmt_words = 0
            self.infmt_letters = 0
            self.loc_words = 0
            self.loc_letters = 0
            self.loc_infmts = 0
            self.prev_infmt = None

        def add_letter(self, infmt, wc):
            self.infmt_words += wc
            self.infmt_letters += 1
            self.loc_words += wc
            self.loc_letters += 1
            if infmt != self.prev_infmt:
                self.loc_infmts += 1
                self.prev_infmt = infmt

        def get_and_reset_informant(self):
            words, letters = self.infmt_words, self.infmt_letters
            self.infmt_words = self.infmt_letters = 0
            return (words, letters)

        def get_and_reset_locality(self):
            words, letters, infmts = (self.loc_words, self.loc_letters,
                                      self.loc_infmts)
            self.loc_words = self.loc_letters = self.loc_infmts = 0
            return (words, letters, infmts)

    def _generate_output(self):
        self._output_data = []
        prev = self.OutputLetterData()
        stats = self.LetterStats()
        self._output_data.append(
            ('file_heading',
             ('Locality, filename and word count in ScotsCorr letters by male,'
              ' female and royal writers in the four time-periods')))
        for filename in self._letter_order:
            if filename not in self._letter_info:
                self.warn('No information found for letter ' + filename)
                continue
            letter_info = self._letter_info[filename]
            try:
                data = self.OutputLetterData(letter_info)
            except WordCountDocMakerError as e:
                self.error(e.args[0], show_filename=False)
            if prev.informant is not None:
                if data.informant != prev.informant:
                    self._output_data.append(
                        ('informant_total', stats.get_and_reset_informant()))
                if data.locality != prev.locality:
                    self._output_data.append(
                        ('locality_total', stats.get_and_reset_locality()))
            if data.gender != prev.gender:
                self._output_data.append(
                    ('gender_heading', data.gender.title() + ' writers'))
            if ((data.gender != prev.gender or data.period != prev.period)
                and data.period != 'royal'):
                self._output_data.append(('period_heading', data.period))
            word_count = letter_info['wc_new']
            stats.add_letter(data.informant, int(word_count))
            self._output_data.append(
                ('letter_info', (data.locality, letter_info['lclet'],
                                 filename, word_count)))
            prev = data
        self._output_data.append(
            ('informant_total', stats.get_and_reset_informant()))
        self._output_data.append(
            ('locality_total', stats.get_and_reset_locality()))

    def _write_output(self):
        writer_class = getattr(
            self, 'OutputWriter' + self._opts.output_format.title())
        writer = writer_class(write_fn=self.output)
        writer.output(self._output_data)

    class OutputWriter(object):

        _tags = {
            'none': ('', ''),
            'bold': ('<b>', '</b>'),
        }

        def __init__(self, write_fn=None):
            self._write_fn = write_fn or self._default_write_fn

        def _default_write_fn(self, text):
            sys.stdout.write(text)

        def output(self, output_data):
            for itemnr, data_item in enumerate(output_data):
                data_type, data = data_item
                try:
                    next_type = output_data[itemnr + 1][0]
                except IndexError:
                    next_type = None
                getattr(self, '_output_' + data_type)(data, next_type)

        def _output_heading(self, data, level):
            self._write_fn(self._make_heading(data, level))

        def _make_heading(self, data, level):
            tagname = u'h' + unicode(level)
            return '<' + tagname + '>' + data + '</' + tagname + '>\n\n'

        def _output_file_heading(self, data, _):
            self._output_heading(data, 1)

        def _output_gender_heading(self, data, _):
            self._output_heading(data, 2)

        def _output_period_heading(self, data, _):
            self._output_heading(data, 3)

        def _output_letter_info(self, data, next_type):
            self._write_fn(self._make_letter_info(data, next_type))

        def _make_letter_info(self, data, next_type):
            locality_inf, locality_letter, filename, word_count = data
            if next_type != 'letter_info':
                mo = re.match(r'(\S+?)(\d{7}.*)$', filename)
                filename = (self._tags['bold'][0] + mo.group(1)
                            + self._tags['bold'][1] + mo.group(2))
            result = ('%LC: {lcinf}, {lclet}\n%FN: {fn}\n%WC: {wc}'
                      .format(lcinf=locality_inf,
                              lclet=locality_letter,
                              fn=filename,
                              wc=word_count))
            if next_type == 'letter_info':
                result += '\n\n'
            return result

        def _output_informant_total(self, data, next_type):
            self._write_fn(self._make_informant_total(data, next_type))

        def _make_informant_total(self, data, next_type):
            words, letters = data
            result = '\t\t\t\t\t{wc}\tLetters {lc}'.format(wc=words, lc=letters)
            if next_type == 'letter_info':
                result += '\n\n'
            return result

        def _output_locality_total(self, data, _):
            self._write_fn(self._make_locality_total(data))

        def _make_locality_total(self, data):
            words, letters, informants = data
            return ('\tTotal {wc}\n\t\t\t\t\t\t\t\t\t{ic}/{lc}\n\n'
                    .format(wc=words, lc=letters, ic=informants))

    class OutputWriterText(OutputWriter):

        pass

    class OutputWriterDocx(OutputWriter):

        def __init__(self, write_fn=None):

            def reset_font(style):
                font = style.font
                font.color.rgb = None
                font.name = 'Times New Roman'
                if font.size < Pt(12):
                    font.size = Pt(12)

            def move_space_after(style):
                parfmt = style.paragraph_format
                if parfmt.space_before > 0:
                    parfmt.space_after = parfmt.space_before
                    parfmt.space_before = 0

            self._write_fn = self._default_write_fn
            self._write_final_fn = write_fn or sys.stdout.write
            self._doc = docx.Document()
            styles = self._doc.styles
            for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Normal']:
                reset_font(styles[style_name])
                move_space_after(styles[style_name])
            styles['Heading 1'].font.small_caps = True
            sect = self._doc.sections[0]
            sect.page_height = Mm(297)
            sect.page_width = Mm(210)
            sect.left_margin = sect.right_margin = sect.top_margin

        def _default_write_fn(self, text):
            import sys
            # sys.stderr.write(repr(text) + '\n')
            text = text.rstrip('\n')
            self._curr_para = self._doc.add_paragraph(text)

        def _output_heading(self, data, level):
            self._doc.add_heading(data, level)

        def _output_letter_info(self, data, next_type):
            text = self._make_letter_info(data, next_type)
            text.rstrip('\n')
            parts = re.split(
                self._tags['bold'][0] + '|' + self._tags['bold'][1], text)
            # sys.stderr.write(repr(parts) + '\n')
            self._write_fn(parts[0])
            for partnr, part in enumerate(parts[1:]):
                run = self._curr_para.add_run(part)
                if partnr % 2 == 0:
                    run.bold = True

        def output(self, output_data):
            super(WordCountDocMaker.OutputWriterDocx, self).output(output_data)
            outf = StringIO()
            self._doc.save(outf)
            self._write_final_fn(outf.getvalue())

        def _output_informant_total(self, data, next_type):
            self._curr_para.add_run(
                self._make_informant_total(data, next_type).rstrip('\n'))

        def _output_locality_total(self, data, _):
            self._curr_para.add_run(self._make_locality_total(data).rstrip('\n'))
            
    def getopts(self, args=None):
        self.getopts_basic(
            dict(usage="%prog [options] [input] > output",
                 description=(
"""Generate a list of letter word counts by gender, period and
locality based on an input file in TSV format containing information
on the variables for each letter.""")
             ),
            args,
            ['order-file = FILE', dict(
                help=('read the order of letters from FILE; FILE should be'
                      ' a plain-text file with lines beginning with "%FN:"'
                      ' marking file names'))],
            ['output-format = FORMAT', dict(
                type='choice',
                choices=['text', 'docx'],
                default='text',
                help='output in FORMAT: text or docx (default: %default)')]
        )
        if self._opts.output_format == 'docx':
            self._output_encoding = None


if __name__ == "__main__":
    WordCountDocMaker().run()
