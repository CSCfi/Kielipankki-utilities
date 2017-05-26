#! /usr/bin/env python
# -*- coding: utf-8 -*-


""" """


import sys
import re

from cStringIO import StringIO

import pandas as pd
from pandas.core.indexing import IndexingError

# Package python-docx: https://python-docx.readthedocs.io/en/latest/
import docx
from docx.shared import Pt, Mm

import korpimport.util as korputil
import korpimport.xmlutil as xmlutil


class DictTable(list):

    def __init__(self):
        super(DictTable, self).__init__()

    def add_row(self, row=None):
        self.append(row or {})

    def add_cell(self, key, value):
        if not len(self):
            self.append({})
        if not isinstance(value, dict):
            value = {'value': value}
        self[-1][key] = value


class StatsMaker(korputil.InputProcessor):

    _localities = [
        ('North', [
            'Moray',
            'Invernessshire',
            'Sutherland',
            'Ross',]),
        ('North-East', [
            'Aberdeenshire',
            'Angus',]),
        ('Central', [
            'Perthshire',
            'Lanarkshire',]),
        ('South-East', [
            'Fife',
            'Lothian',
            'Borders',
            'Stirlingshire',]),
        ('South-West', [
            'Ayrshire',
            'Argyllshire',
            'South-West',]),
        ('Professional', ['Professional']),
        ('Unlocalised', ['unlocalised']),
    ]
    _periods = [u'1540–1599', u'1600–1649', u'1650–1699', u'1700–1749']
    _genders = [u'male', u'female', u'royal']

    def __init__(self):
        super(StatsMaker, self).__init__()
        self._letter_info = None
        self._output_tables = []
        self._locality_order = [(loc, reg) for reg, locs in self._localities
                                for loc in locs]
        writer_class = getattr(
            self, 'OutputWriter' + self._opts.output_format.title())
        self._writer = writer_class(write_fn=self.output)

    def process_input_stream(self, stream, filename=None):
        self._read_wordcount_file(stream)
        self._generate_output()
        self._write_output()

    def _read_wordcount_file(self, stream):
        self._letter_info = pd.read_csv(
            stream, sep='\t', encoding=self._input_encoding, index_col='fn')
        self._letter_info = self._letter_info[
            ['period', 'gender', 'lcinf', 'largeregion', 'wc_new', 'from']]
        self._letter_info = self._letter_info.rename(columns={'wc_new': 'wc'})

    def _generate_output(self):
        make_table_fns = [(self._make_by_locality, True),
                          (self._make_overview, False),
                          (self._make_by_larger_region, True),
                          (self._make_by_category, False),
                          (self._make_geodef, False),
                          (self._make_informants_by_locality, True),
                          (self._make_by_category_period, False)]
        for make_table_fn, split_by_gender in make_table_fns:
            self._add_output_tables(make_table_fn, split_by_gender)

    def _add_output_tables(self, make_table_fn, split_by_gender=False):
        if split_by_gender:
            for gender in ['male', 'female']:
                self._output_tables.append(make_table_fn(gender))
        else:
            self._output_tables.append(make_table_fn())

    def _make_by_locality(self, gender):
        table = DictTable()
        letters = self._letter_info.query(
            'gender == "{gender}"'.format(gender=gender))
        counts = self._make_counts(letters, ['lcinf', 'period'])
        loc_counts = self._make_counts(letters, ['lcinf'])
        lr_counts = self._make_counts(letters, ['largeregion'])
        total_counts = self._make_counts(letters, ['period'])
        token_count_total_all = lr_counts['tokens'].sum()
        token_count_total_geodef = (
            letters.query('lcinf != "Professional" and lcinf != "unlocalised"')
            ['wc'].sum())
        prev_lr = None

        def add_larger_region_totals(table, region):
            table[-1]['Larger region']['bold'] = True
            table.add_cell('Larger region total',
                           self._get_counts(lr_counts, region))
            table.add_cell(
                'Larger region total %',
                (lr_counts['tokens'][region] * 1.0
                 / token_count_total_all * 100.0))

        def add_totals(table):
            table.add_row()
            table.add_cell('Locality', 'Total')
            table.add_cell('Larger region', '')
            cnts = dict((key, loc_counts[key].sum())
                        for key in ['informants', 'letters', 'tokens'])
            table.add_cell('N Writers', cnts['informants'])
            table.add_cell('N Letters', cnts['letters'])
            table.add_cell('Total', cnts['tokens'])
            table.add_cell('Larger region total', cnts)
            table.add_cell('Larger region total %', 100.0)
            for period in self._periods:
                table.add_cell(period, self._get_counts(total_counts, period))
            table.add_row()
            table.add_cell('Locality', 'Total %')
            for period in self._periods:
                table.add_cell(
                    period,
                    (self._get_counts(total_counts, period)['tokens'] * 1.0
                     / token_count_total_all * 100.0))
            table.add_cell('Larger region total %', 100.0)

        for locality, larger_region in self._locality_order:
            if not loc_counts['informants'].get(locality):
                continue
            if larger_region != prev_lr and prev_lr is not None:
                add_larger_region_totals(table, prev_lr)
            table.add_row()
            table.add_cell('Locality', locality)
            table.add_cell('Larger region', larger_region)
            # print repr(counts['informants']).decode('utf-8')
            table.add_cell('N Writers', loc_counts['informants'].ix[locality])
            table.add_cell('N Letters', loc_counts['letters'].ix[locality])
            token_cnt = loc_counts['tokens'].ix[locality]
            table.add_cell('Total', token_cnt)
            if locality not in ['unlocalised', 'Professional']:
                table.add_cell(
                    '% of geographically defined',
                    token_cnt * 1.0 / token_count_total_geodef * 100.0)
            for period in self._periods:
                table.add_cell(period,
                               self._get_counts(counts, (locality, period)))
            prev_lr = larger_region
        add_larger_region_totals(table, prev_lr)
        add_totals(table)
        # print repr(table).decode('utf-8')
        formatted_table = self._format_table(
            table, (['Locality', 'N Writers'] + self._periods
                    + ['Total', 'N Letters', '% of geographically defined',
                       'Larger region', 'Larger region total',
                       'Larger region total %']),
            headings='bold')
        return {
            'table': formatted_table,
            'title': gender.title() + (u' informants in the Helsinki Corpus of'
                                       u' Scottish Correspondence 1540–1750'),
        }

    def _make_counts(self, df, group_cols):
        counts = {}
        groups = df.groupby(group_cols)
        counts['letters'] = groups.count()['wc']
        counts['tokens'] = groups['wc'].sum()
        # Hint for this from
        # http://stackoverflow.com/questions/17679089/pandas-dataframe-groupby-two-columns-and-get-counts
        counts['informants'] = (
            groups['from']
            .value_counts()
            .groupby(level=range(len(group_cols)))
            .count())
        return counts

    def _get_counts(self, counts, key):
        # print key
        try:
            # for count_type in counts.iterkeys():
            #     print count_type, key, type(counts[count_type]), repr(counts[count_type].ix[key] if isinstance(key, basestring) else counts[count_type].ix[key[0]][key[1]]).decode('utf-8')
            return dict([(count_type, counts[count_type].ix[key])
                         for count_type in counts.iterkeys()])
        except (KeyError, IndexingError):
            # print "KeyError"
            return {}

    def _make_overview(self):
        table = DictTable()
        counts = {}
        counts['tokens'] = (
            self._letter_info.groupby(['period', 'gender'])['wc'].sum())
        tokens_total = counts['tokens'].sum()
        # sys.stdout.write(repr(counts['tokens']).decode('utf-8'))
        counts['period_totals'] = (
            self._letter_info.groupby(['period'])['wc'].sum())
        counts['gender_totals'] = (
            self._letter_info.groupby(['gender'])['wc'].sum())
        counts['informants'] = (
            self._letter_info.groupby(['gender'])
            ['from']
            .value_counts()
            .groupby(level=0)
            .count())
        counts['letters'] = self._letter_info.groupby(['gender']).count()['wc']
        for gender in self._genders:
            table.add_row()
            table.add_cell('Gender/rank',
                           gender.title() if gender != 'royal' else 'Court')
            for period in self._periods:
                try:
                    count = counts['tokens'].ix[period].ix[gender]
                except (KeyError, IndexingError):
                    count = {}
                table.add_cell(period, count)
            table.add_cell('Total', counts['gender_totals'].ix[gender])
            table.add_cell(
                '%', 100.0 * counts['gender_totals'].ix[gender] / tokens_total)
            table.add_cell('N Informants', counts['informants'].ix[gender])
            table.add_cell('N Letters', counts['letters'].ix[gender])
        table.add_row()
        table.add_cell('Gender/rank', 'Total')
        for period in self._periods:
            table.add_cell(period, counts['period_totals'].ix[period])
        table.add_cell('Total', tokens_total)
        table.add_cell('N Informants', counts['informants'].sum())
        table.add_cell('N Letters', counts['letters'].sum())
        table.add_row()
        table.add_cell('Gender/rank', '%')
        for period in self._periods:
            table.add_cell(
                period,
                100.0 * counts['period_totals'].ix[period] / tokens_total)
        table.add_cell('%', 100.0)
        formatted_table = self._format_table(
            table, (['Gender/rank'] + self._periods
                    + ['Total', '%', 'N Informants', 'N Letters']),
            headings='bold')
        return {
            'table': formatted_table,
            'title': (u'Informants in the Helsinki Corpus of'
                      u' Scottish Correspondence 1540–1750'),
        }

    def _make_by_larger_region(self, gender):
        table = DictTable()
        letters = self._letter_info.query(
            'gender == "{gender}"'.format(gender=gender))
        lr_counts = self._make_counts(letters, ['largeregion'])['tokens']
        # sys.stdout.write(repr(lr_counts))
        token_count_total_geodef = (
            letters.query('lcinf != "Professional" and lcinf != "unlocalised"')
            ['wc'].sum())
        # print(token_count_total_geodef)
        larger_regions = [lr for lr, regs in self._localities
                          if lr not in ['Professional', 'Unlocalised']]
        for larger_region in larger_regions:
            table.add_row()
            table.add_cell('Area', larger_region)
            table.add_cell(
                'Percentage',
                100.0 * lr_counts.ix[larger_region] / token_count_total_geodef)
        table.add_row()
        table.add_cell('Area', 'Total')
        table.add_cell('Percentage', 100.0)
        formatted_table = self._format_table(table, ['Area', 'Percentage'],
                                             headings='bold')
        return {
            'table': formatted_table,
            'title': (u'Percentages of correspondence by ' + gender
                      + u' writers in the five larger areas of Scotland in'
                      + u'the ScotsCorr corpus'),
        }

    def _make_geodef(self):
        table = DictTable()
        letters = self._letter_info.query(
            'lcinf != "Professional" and lcinf != "unlocalised"'
            ' and lcinf != "Court"')
        counts = letters.groupby(['largeregion', 'gender'])['wc'].sum()
        counts_gender = letters.groupby(['gender'])['wc'].sum()
        token_count_total_geodef = letters['wc'].sum()
        larger_regions = [lr for lr, regs in self._localities
                          if lr not in ['Professional', 'Unlocalised']]
        for larger_region in larger_regions:
            table.add_row()
            table.add_cell('Region', larger_region)
            for gender in ['male', 'female']:
                table.add_cell(gender.title(), counts.ix[larger_region, gender])
            total = counts.ix[larger_region].sum()
            table.add_cell('Total', total)
            table.add_cell('%', 100.0 * total / token_count_total_geodef)
        table.add_row()
        table.add_cell('Region', 'Total')
        for gender in ['male', 'female']:
            table.add_cell(gender.title(), counts_gender[gender])
            table.add_cell('Total', token_count_total_geodef)
            table.add_cell('%', 100.0)
        formatted_table = self._format_table(
            table, ['Region', 'Male', 'Female', 'Total', '%'],
            headings='bold')
        return {
            'table': formatted_table,
            'title': (u'Geographically defined informants in the ScotsCorr'
                      u' corpus'),
        }

    def _make_by_category(self):
        table = DictTable()
        total = self._letter_info['wc'].sum()
        table.add_row()
        table.add_cell('Category', 'Geographically defined')
        token_count = self._letter_info.query(
            'lcinf != "Professional" and lcinf != "unlocalised"'
            ' and lcinf != "Court"')['wc'].sum()
        table.add_cell('Words', token_count)
        table.add_cell('%', 100.0 * token_count / total)
        for lcinf in ['Professional', 'unlocalised', 'Court']:
            token_count = (
                self._letter_info.query('lcinf == "{0}"'.format(lcinf))
                ['wc'].sum())
            table.add_row()
            table.add_cell('Category', lcinf.title())
            table.add_cell('Words', token_count)
            table.add_cell('%', 100.0 * token_count / total)
        table.add_row()
        table.add_cell('Category', 'Total')
        table.add_cell('Words', total)
        table.add_cell('%', 100.0)
        formatted_table = self._format_table(
            table, ['Category', 'Words', '%'], headings='bold')
        return {
            'table': formatted_table,
            'title': (u'Number of words by letter category in the ScotsCorr'
                      u' corpus'),
        }

    def _make_informants_by_locality(self, gender):
        table = DictTable()
        letters = self._letter_info.query(
            'gender == "{gender}"'.format(gender=gender))
        loc_counts = self._make_counts(letters, ['lcinf'])
        lr_counts = self._make_counts(letters, ['largeregion'])
        # total_counts = self._make_counts(letters, [])

        def get_count(count, locality):
            try:
                return count.ix[locality]
            except (KeyError, IndexingError):
                return {}

        def get_count_str(count, locality):
            return unicode(get_count(count, locality) or u'–')

        def add_row(table, locality, counts):
            table.add_row()
            table.add_cell('Locality', locality)
            table.add_cell('Number of informants',
                           get_count(counts['informants'], locality))
            table.add_cell('Number of letters',
                           get_count(counts['letters'], locality))

        for larger_region, localities in self._localities:
            for locality in localities:
                add_row(table, locality, loc_counts)
            table.add_cell(
                'Larger region',
                larger_region + ' '
                + get_count_str(lr_counts['informants'], larger_region)
                + ' / '
                + get_count_str(lr_counts['letters'], larger_region))
        table.add_row()
        table.add_cell('Locality', 'Total')
        table.add_cell('Number of informants',
                       letters['from'].value_counts().count())
        table.add_cell('Number of letters', letters['wc'].count())
        formatted_table = self._format_table(
            table, ['Locality', 'Number of informants', 'Number of letters',
                    'Larger region'],
            headings='bold')
        return {
            'table': formatted_table,
            'title': (u'Number of ' + gender + ' informants and their letters'
                      u' by locality or district in the ScotsCorr corpus'),
        }

    def _make_by_category_period(self):
        cats = ['Male', 'Male Professional', 'Male unlocalised',
                'Female', 'Female unlocalised', 'Court']
        table = DictTable()
        total_words = self._letter_info['wc'].sum()

        def add_row(table, letters, cat):
            table.add_row()
            table.add_cell('Category', cat)
            for period in self._periods:
                table.add_cell(
                    period, letters.query(u'period == u"{period}"'
                                          .format(period=period))['wc'].sum())
                period_total = letters['wc'].sum()
                table.add_cell('Total', period_total)
                table.add_cell('%', 100.0 * period_total / total_words)
                table.add_cell('N Letters', letters['wc'].count())
                table.add_cell('N Informants',
                               letters['from'].value_counts().count())

        for cat in cats:
            if ' ' in cat:
                gender, lcinf = cat.split()
                lcinf_cond = ' and lcinf == "{lcinf}"'.format(lcinf=lcinf)
            elif cat == 'Court':
                gender = 'royal'
                lcinf_cond = ''
            else:
                gender = cat
                lcinf_cond = (
                    ' and lcinf != "Court" and lcinf != "Professional"'
                    ' and lcinf != "unlocalised"')
            letters = self._letter_info.query(
                'gender == "{gender}"{lcinf_cond}'.format(
                    gender=gender.lower(), lcinf_cond=lcinf_cond))
            add_row(table, letters, cat)
        add_row(table, self._letter_info, 'Total')
        table.add_row()
        table.add_cell('Category', '%')
        for period in self._periods:
            table.add_cell(period,
                           100.0 * table[-2][period]['value'] / total_words)
            table.add_cell('Total', 100.0)
        formatted_table = self._format_table(
            table, (['Category'] + self._periods
                    + ['Total', '%', 'N Informants', 'N Letters']),
            headings='bold')
        return {
            'table': formatted_table,
            'title': (u'Geographically defined male and female writers,'
                      u' professional male writers, and unlocalized male and'
                      ' female writers in the ScotsCorr corpus'),
        }

    def _format_table(self, table, col_order, headings=None):
        formatted_table = []
        formatted_table.append([{'value': colhead} for colhead in col_order])
        for row in table:
            formatted_row = []
            for key in col_order:
                cell = row.get(key)
                if cell is None:
                    cell = {'value': ''}
                elif 'value' not in cell:
                    if cell.get('tokens'):
                        cell['value'] = (
                            u'{tokens:,d}\n{informants:,d} / {letters:,d}'
                            .format(**dict((key, cell.get(key))
                                           for key in ['tokens', 'informants',
                                                       'letters'])))
                    else:
                        cell['value'] = u'–'
                    cell['class'] = 'num'
                else:
                    value = cell.get('value')
                    if isinstance(value, float):
                        cell['value'] = '{0:.1f}'.format(value)
                    elif isinstance(value, int):
                        cell['value'] = '{0:,d}'.format(value)
                    else:
                        cell['value'] = unicode(value)
                    if isinstance(value, float) or isinstance(value, int):
                        cell['class'] = 'num'
                formatted_row.append(cell)
            formatted_table.append(formatted_row)
        # print repr(formatted_table)
        if headings != None:
            for col in formatted_table[0]:
                col[headings] = True
            for row in formatted_table:
                row[0][headings] = True
        return formatted_table

    def _write_output(self):
        self._writer.output(self._output_tables)
        # print self._output_tables

    class OutputWriter(object):

        _tags = {
            'none': ('', ''),
            'bold': ('<b>', '</b>'),
        }

        def __init__(self, write_fn=None):
            self._write_fn = write_fn or self._default_write_fn

        def _default_write_fn(self, text):
            sys.stdout.write(text)

        def output(self, tables_info):
            self._output_header()
            for tablenum, table_info in enumerate(tables_info):
                self._output_heading('Table ' + unicode(tablenum + 1) + '. '
                                     + table_info['title'], 2)
                self._output_table(table_info['table'])
            self._output_footer()

        def _output_header(self):
            self._write_fn(u'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<style>
th {text-align: left; vertical-align: top;  max-width: 10em; padding: 0.5em;}
td {vertical-align: top; padding-right: 0.5em; padding-bottom: 0.25em;}
td.num {text-align: right;}
</style>
</head>
<body>
''')

        def _output_footer(self):
            self._write_fn(u'</body>\n</html>\n')

        def _output_heading(self, heading, level):
            self._write_fn(self._make_heading(heading, level))

        def _make_heading(self, data, level):
            tagname = u'h' + unicode(level)
            return xmlutil.make_elem(tagname, data) + '\n'

        def _output_table(self, table):
            self._write_fn(self._make_table(table))

        def _make_table(self, table):
            return xmlutil.make_elem('table', self._make_table_rows(table),
                                     newlines=True)

        def _make_table_rows(self, table):
            result = [
                (xmlutil.make_elem('tr', self._make_table_row(row, rownum))
                 + '\n') for rownum, row in enumerate(table)]
            return ''.join(result)

        def _make_table_row(self, row, rownum):
            result = [
                xmlutil.make_elem(
                    'th' if rownum == 0 or colnum == 0 else 'td',
                    self._make_table_cell(cell),
                    attrs=[('class', cell['class'])] if 'class' in cell else [])
                for colnum, cell in enumerate(row)]
            return ''.join(result)

        def _make_table_cell(self, cell):
            value = cell['value'].replace('\n', '<br/>')
            return (xmlutil.make_elem('strong', value) if cell.get('bold')
                    else value)

    class OutputWriterHtml(OutputWriter):

        pass

    class OutputWriterDocx(OutputWriter):

        def __init__(self, write_fn=None):

            def reset_font(style):
                font = style.font
                font.color.rgb = None
                font.name = 'Times New Roman'
                if font.size < Pt(12):
                    font.size = Pt(12)

            def move_space_after(style):
                parfmt = style.paragraph_format
                if parfmt.space_before > 0:
                    parfmt.space_after = parfmt.space_before
                    parfmt.space_before = 0

            self._write_fn = self._default_write_fn
            self._write_final_fn = write_fn or sys.stdout.write
            self._doc = docx.Document()
            styles = self._doc.styles
            for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Normal']:
                reset_font(styles[style_name])
                move_space_after(styles[style_name])
            styles['Heading 1'].font.small_caps = True
            sect = self._doc.sections[0]
            sect.page_height = Mm(297)
            sect.page_width = Mm(210)
            sect.left_margin = sect.right_margin = sect.top_margin

        def _default_write_fn(self, text):
            import sys
            # sys.stderr.write(repr(text) + '\n')
            text = text.rstrip('\n')
            self._curr_para = self._doc.add_paragraph(text)

        def _output_heading(self, data, level):
            self._doc.add_heading(data, level)

        def output(self, output_data):
            super(StatsMaker.OutputWriterDocx, self).output(output_data)
            outf = StringIO()
            self._doc.save(outf)
            self._write_final_fn(outf.getvalue())
            
    def getopts(self, args=None):
        self.getopts_basic(
            dict(usage="%prog [options] [input] > output",
                 description=(
""" """)
             ),
            args,
            ['output-format = FORMAT', dict(
                type='choice',
                choices=['html', 'docx'],
                default='html',
                help='output in FORMAT: html or docx (default: %default)')]
        )
        if self._opts.output_format == 'docx':
            self._output_encoding = None


if __name__ == "__main__":
    StatsMaker().run()
